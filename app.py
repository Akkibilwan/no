import streamlit as st
from docx import Document
from docx.shared import Inches

def generate_biodata():
    """Generates a biodata document."""

    # Create a new Word document
    doc = Document()

    # Title
    doc.add_heading("Biodata for Marriage", 0)

    # Personal Information
    st.subheader("Personal Information")
    doc.add_heading("Personal Information", level=1)

    full_name = st.text_input("Full Name")
    if full_name:
        doc.add_paragraph(f"Full Name: {full_name}")

    date_of_birth = st.date_input("Date of Birth")
    if date_of_birth:
        doc.add_paragraph(f"Date of Birth: {date_of_birth.strftime('%Y-%m-%d')}")

    age = st.number_input("Age", min_value=18, max_value=100, step=1)
    if age:
        doc.add_paragraph(f"Age: {age}")

    height = st.number_input("Height (in cm)", min_value=100, max_value=250, step=1)
    if height:
        doc.add_paragraph(f"Height: {height} cm")

    weight = st.number_input("Weight (in kg)", min_value=30, max_value=200, step=1)
    if weight:
        doc.add_paragraph(f"Weight: {weight} kg")

    complexion = st.selectbox("Complexion", ["Fair", "Wheatish", "Dark", "Other"])
    if complexion:
        doc.add_paragraph(f"Complexion: {complexion}")

    blood_group = st.selectbox("Blood Group", ["A+", "A-", "B+", "B-", "AB+", "AB-", "O+", "O-"])
    if blood_group:
        doc.add_paragraph(f"Blood Group: {blood_group}")

    health_status = st.text_area("Health Status (Optional)")
    if health_status:
        doc.add_paragraph(f"Health Status: {health_status}")

    # Family Details
    st.subheader("Family Details")
    doc.add_heading("Family Details", level=1)

    father_name = st.text_input("Father's Name")
    if father_name:
        doc.add_paragraph(f"Father's Name: {father_name}")

    mother_name = st.text_input("Mother's Name")
    if mother_name:
        doc.add_paragraph(f"Mother's Name: {mother_name}")

    siblings = st.text_area("Siblings (Names, Ages, Occupations)")
    if siblings:
        doc.add_paragraph(f"Siblings: {siblings}")

    family_background = st.text_area("Family Background (Brief description, including profession and status)")
    if family_background:
        doc.add_paragraph(f"Family Background: {family_background}")

    # Contact Information
    st.subheader("Contact Information")
    doc.add_heading("Contact Information", level=1)

    permanent_address = st.text_area("Permanent Address")
    if permanent_address:
        doc.add_paragraph(f"Permanent Address: {permanent_address}")

    current_address = st.text_area("Current Address (if different from permanent)")
    if current_address:
        doc.add_paragraph(f"Current Address: {current_address}")

    phone_number = st.text_input("Phone Number")
    if phone_number:
        doc.add_paragraph(f"Phone Number: {phone_number}")

    email_address = st.text_input("Email Address")
    if email_address:
        doc.add_paragraph(f"Email Address: {email_address}")

    # Educational Background
    st.subheader("Educational Background")
    doc.add_heading("Educational Background", level=1)

    highest_qualification = st.text_input("Highest Qualification")
    if highest_qualification:
        doc.add_paragraph(f"Highest Qualification: {highest_qualification}")

    college_university = st.text_input("College/University Attended")
    if college_university:
        doc.add_paragraph(f"College/University Attended: {college_university}")

    other_qualifications = st.text_area("Other Qualifications or Certifications")
    if other_qualifications:
        doc.add_paragraph(f"Other Qualifications or Certifications: {other_qualifications}")

    # Professional Details
    st.subheader("Professional Details")
    doc.add_heading("Professional Details", level=1)

    current_occupation = st.text_input("Current Occupation")
    if current_occupation:
        doc.add_paragraph(f"Current Occupation: {current_occupation}")

    job_title = st.text_input("Job Title")
    if job_title:
        doc.add_paragraph(f"Job Title: {job_title}")

    company_organization = st.text_input("Company/Organization")
    if company_organization:
        doc.add_paragraph(f"Company/Organization: {company_organization}")

    work_experience = st.text_area("Work Experience (Brief Summary)")
    if work_experience:
        doc.add_paragraph(f"Work Experience: {work_experience}")

    annual_income = st.number_input("Annual Income")
    if annual_income:
        doc.add_paragraph(f"Annual Income: {annual_income}")

    # Lifestyle and Interests
    st.subheader("Lifestyle and Interests")
    doc.add_heading("Lifestyle and Interests", level=1)

    dietary_preferences = st.selectbox("Dietary Preferences", ["Vegetarian", "Non-Vegetarian", "Vegan", "Other"])
    if dietary_preferences:
        doc.add_paragraph(f"Dietary Preferences: {dietary_preferences}")

    hobbies_interests = st.text_area("Hobbies and Interests")
    if hobbies_interests:
        doc.add_paragraph(f"Hobbies and Interests: {hobbies_interests}")

    languages_known = st.text_area("Languages Known")
    if languages_known:
        doc.add_paragraph(f"Languages Known: {languages_known}")

    # Download the document
    if st.button("Generate Biodata"):
        st.download_button(
            label="Download Biodata",
            data=doc.save(f"{full_name}_Biodata.docx"),
            file_name=f"{full_name}_Biodata.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

if __name__ == "__main__":
    st.set_page_config(page_title="Biodata Generator")
    st.title("Biodata Generator for Marriage")
    generate_biodata()
